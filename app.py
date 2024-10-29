import os
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from models import db, User, Category, Product, Order, OrderItem, Cart, Manager, ChatMessage
from forms import LoginForm, RegistrationForm, UpdateProfileForm, AddProductForm, AddCategoryForm, EditProductForm, EditCategoryForm, OrderForm
from flask_login import LoginManager, current_user, login_user, logout_user, login_required
from werkzeug.utils import secure_filename
from sqlalchemy import func
from datetime import datetime
from flask_admin import Admin, expose
from flask_admin.contrib.sqla import ModelView
from flask_admin.menu import MenuLink
from flask_admin.base import AdminIndexView
import telebot
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+mysqlconnector://admin:admin@localhost/ooo_btt_db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'your_secret_key' 

BOT_TOKEN = '7092043703:AAFCY9XIcUK5km_EMy3FfYUEHccxT8XN0_8' 
bot = telebot.TeleBot(BOT_TOKEN) 

UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

db.init_app(app)
migrate = Migrate(app, db)

login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message_category = 'info'


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


from functools import wraps


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('У вас нет прав доступа  к этой странице.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


def manager_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_manager:
            flash('У вас нет прав доступа  к этой странице.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


def calculate_order_total(order_id):
    order = Order.query.get(order_id)
    if not order:
        return

    order.total_price = db.session.query(
        func.sum(OrderItem.quantity * Product.price)
    ).join(Product).filter(OrderItem.order_id == order_id).scalar()

    db.session.commit()


@app.context_processor
def inject_categories():
    """Делает категории доступными во всех шаблонах."""
    categories = Category.query.all()  # Получаем все категории
    return dict(categories=categories)



# === Представления Flask-Admin ===

class MyModelView(ModelView):
    def is_accessible(self):
        return current_user.is_authenticated and current_user.is_admin

    def can_create(self):
        return False

    def can_delete(self):
        return False

class MyAdminIndexView(AdminIndexView):
    @expose('/')
    def index(self):
        return self.render('admin/index.html', admin=admin)

# Создаем экземпляр Flask-Admin
admin = Admin(app, name='Админ-панель', 
              template_mode='bootstrap4', index_view=MyAdminIndexView())

# Добавляем модели в админ-панель
admin.add_view(MyModelView(User, db.session))
admin.add_view(MyModelView(Category, db.session))
admin.add_view(MyModelView(Product, db.session))
admin.add_view(MyModelView(Order, db.session))
admin.add_view(MyModelView(OrderItem, db.session))
admin.add_view(MyModelView(Cart, db.session))
admin.add_view(MyModelView(Manager, db.session))

# Добавляем ссылку для выхода в меню админки
admin.add_link(MenuLink(name='Выход', url='/', endpoint='admin.index')) 

# ===  Flask маршруты ===

@app.route('/')
def index():
    products_query = Product.query
    min_price = db.session.query(func.min(Product.price)).scalar()
    max_price = db.session.query(func.max(Product.price)).scalar()

    # Фильтрация по цене
    price_min = request.args.get('price_min')
    price_max = request.args.get('price_max')
    if price_min:
        products_query = products_query.filter(Product.price >= price_min)
    if price_max:
        products_query = products_query.filter(Product.price <= price_max)

    # Фильтрация избранного
    if request.args.get('favorites_only') and current_user.is_authenticated:
        products_query = products_query.filter(Product.id.in_([fav.id for fav in current_user.favorites]))

    products = products_query.all()
    return render_template('index.html', products=products, min_price=min_price, max_price=max_price)

@app.route('/favorite/<int:product_id>')
@login_required
def favorite(product_id):
    product = Product.query.get_or_404(product_id)
    if product in current_user.favorites:
        current_user.favorites.remove(product)
        db.session.commit()
        flash('Товар удален из избранного.', 'info')
    else:
        current_user.favorites.append(product)
        db.session.commit()
        flash('Товар добавлен в избранное.', 'success')
    return redirect(url_for('index'))  # Или на страницу, где был товар


@app.route('/category/<int:category_id>')
def products_by_category(category_id):
    category = Category.query.get_or_404(category_id)
    products = Product.query.filter_by(category_id=category_id).all()
    return render_template('category.html', category=category, products=products)


@app.route('/product/<int:product_id>')
def product_details(product_id):
    product = Product.query.get_or_404(product_id)

    return render_template(
        'product_details.html',
        product=product,
    )



@app.route('/cart')
def cart():
    if current_user.is_authenticated:
        cart_items = Cart.query.filter_by(user_id=current_user.id).all()
    else:
        cart_items = session.get('cart', [])

    total_price = 0

    for cart_item in cart_items:
        if current_user.is_authenticated:
            product = cart_item.product
            total_price += product.price * cart_item.quantity
              
        else:
            product = Product.query.get(cart_item['product_id'])
            if product:
                cart_item['product_name'] = product.name
                cart_item['product_price'] = product.price
                total_price += product.price * cart_item['quantity']
               

    return render_template('cart.html', cart_items=cart_items, total_price=total_price)  # Передаем обе цены в шаблон

@app.route('/add_to_cart/<int:product_id>')
def add_to_cart(product_id):
    if current_user.is_authenticated:
        cart_item = Cart.query.filter_by(user_id=current_user.id, product_id=product_id).first()
        if cart_item:
            cart_item.quantity += 1
        else:
            cart_item = Cart(user_id=current_user.id, product_id=product_id, quantity=1)
            db.session.add(cart_item)
        db.session.commit()
    else:
        if 'cart' not in session:
            session['cart'] = []
        existing_item = next(
            (item for item in session['cart'] if item['product_id'] == product_id), None)
        if existing_item:
            existing_item['quantity'] += 1
        else:
            session['cart'].append({'product_id': product_id, 'quantity': 1})
        session.modified = True
    return redirect(url_for('cart'))


@app.route('/update_cart', methods=['POST'])
def update_cart():
    if current_user.is_authenticated:
        for key, value in request.form.items():
            if key.startswith('quantity_'):
                try:
                    cart_item_id = int(key.replace('quantity_', ''))
                    quantity = int(value)
                    cart_item = Cart.query.get(cart_item_id)
                    if cart_item:
                        if quantity > 0:
                            cart_item.quantity = quantity
                            db.session.commit()
                        else:
                            db.session.delete(cart_item)
                            db.session.commit()
                except (ValueError, TypeError):
                    pass
    else:
        if 'cart' in session:
            for i in range(len(session['cart'])):
                product_id = session['cart'][i]['product_id']
                quantity_key = f'quantity_{product_id}'
                if quantity_key in request.form:
                    try:
                        new_quantity = int(request.form[quantity_key])
                        if new_quantity > 0:
                            session['cart'][i]['quantity'] = new_quantity
                        else:
                            del session['cart'][i]
                    except (ValueError, TypeError):
                        pass
            session.modified = True
    return redirect(url_for('cart'))


@app.route('/delete_from_cart/<int:item_id>')
def delete_from_cart(item_id):
    if current_user.is_authenticated:
        cart_item = Cart.query.get_or_404(item_id)
        if cart_item.user_id != current_user.id:
            flash('У вас нет прав для удаления этого товара из корзины.', 'danger')
            return redirect(url_for('cart'))

        db.session.delete(cart_item)
        db.session.commit()
        flash('Товар удален из корзины.', 'success')
    else:
        if 'cart' in session:
            session['cart'] = [item for item in session['cart'] if item['product_id'] != item_id]
            session.modified = True
            flash('Товар удален из корзины.', 'success')

    return redirect(url_for('cart'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user is None or not user.check_password(form.password.data):
            flash('Неверный email или пароль', 'danger')
            return redirect(url_for('login'))
        login_user(user, remember=form.remember_me.data)

        # Перенести товары из сессии в базу данных после авторизации
        if 'cart' in session:
            for item in session['cart']:
                # Проверяем, есть ли уже такой товар в корзине пользователя
                existing_cart_item = Cart.query.filter_by(user_id=user.id, product_id=item['product_id']).first()
                if existing_cart_item:
                    # Если товар уже есть, увеличиваем количество
                    existing_cart_item.quantity += item['quantity']
                else:
                    # Если товара нет, добавляем новую запись в корзину
                    cart_item = Cart(user_id=user.id, product_id=item['product_id'], quantity=item['quantity'])
                    db.session.add(cart_item)
            db.session.commit()
            session['cart'] = []  # Очищаем корзину в сессии

        flash('Вы успешно вошли в систему', 'success')
        next_page = request.args.get('next')
        return redirect(next_page) if next_page else redirect(url_for('index'))
    return render_template('login.html', title='Авторизация', form=form)


@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = RegistrationForm()
    if form.validate_on_submit():
        user = User(email=form.email.data, first_name=form.first_name.data,
                    last_name=form.last_name.data, phone=form.phone.data)
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('Поздравляем с успешной регистрацией!', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', title='Регистрация', form=form)


@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    form = UpdateProfileForm()
    if form.validate_on_submit():
        current_user.email = form.email.data
        current_user.first_name = form.first_name.data
        current_user.last_name = form.last_name.data
        current_user.phone = form.phone.data
        db.session.commit()
        flash('Данные профиля успешно обновлены.', 'success')
        return redirect(url_for('profile'))
    elif request.method == 'GET':
        form.email.data = current_user.email
        form.first_name.data = current_user.first_name
        form.last_name.data = current_user.last_name
        form.phone.data = current_user.phone
    return render_template('profile.html', title='Профиль', form=form)


@app.route('/orders')
@login_required
def orders():
    user_orders = Order.query.filter_by(user_id=current_user.id).all()
    return render_template('orders.html', orders=user_orders)



def send_order_notification(order_id):
    order = Order.query.get_or_404(order_id)
    user = order.user
    items_text = '\n'.join([
        f'- {item.product.name} x{item.quantity} = {item.quantity * item.unit_price:.2f} руб.' 
        for item in order.items
    ])

    message = f"Новый заказ!\n" \
              f"Номер заказа: {order.id}\n" \
              f"Дата: {order.order_date.strftime('%Y-%m-%d %H:%M')}\n" \
              f"Пользователь: {user.first_name} {user.last_name} ({user.email})\n" \
              f"Адрес доставки: {order.address}\n" \
              f"Способ оплаты: {order.payment_method}\n\n" \
              f"Товары:\n{items_text}\n\n" \
              f"Итого: {order.total_price:.2f} руб."

    bot.send_message(712388013, message) 
 

@app.route('/checkout', methods=['GET', 'POST'])
@login_required
def checkout():
    cart_items = Cart.query.filter_by(user_id=current_user.id).all()
    form = OrderForm()
    total_price = 0

    if cart_items:
        if request.method == 'POST':
            if form.validate_on_submit():
                # Получение данных из формы
                address = form.address.data
                payment_method = form.payment_method.data
                
                # Создание заказа
                order = Order(
                    user_id=current_user.id, 
                    order_date=datetime.now(), 
                    status='Новый', 
                    address=address,  # Сохранение адреса
                    payment_method=payment_method  # Сохранение способа оплаты
                )
                db.session.add(order)
                db.session.flush()  # Необходимо для получения id заказа

                # Добавление товаров к заказу
                for cart_item in cart_items:
                    product = Product.query.get(cart_item.product_id)
                    if product:
                        price = product.price
                        order_item = OrderItem(order_id=order.id, product_id=cart_item.product_id, quantity=cart_item.quantity, unit_price=price)
                        total_price += price * cart_item.quantity
                        db.session.add(order_item)

                order.total_price = total_price  # Сохранение итоговой суммы
                db.session.commit()


                try:
                    send_order_notification(order.id) 
                except Exception as e:
                    print(f"Ошибка отправки уведомления в Telegram: {e}")

                # Очистка корзины после оформления заказа
                for cart_item in cart_items:
                    db.session.delete(cart_item)
                db.session.commit()
                flash('Заказ успешно оформлен!', 'success')
                return redirect(url_for('orders'))  # Redirect to orders page
            else:
                for field, errors in form.errors.items():
                    flash(f"Ошибка в поле {field}: {errors[0]}", 'danger')
        
        
        for cart_item in cart_items:
            product = cart_item.product
            total_price += product.price * cart_item.quantity
    else:
        flash('Ваша корзина пуста.', 'warning')
        return redirect(url_for('index'))
    return render_template('checkout.html', cart_items=cart_items, form=form, total_price=total_price)


@app.route('/create_order', methods=['POST'])
@login_required
def create_order():
    cart_items = Cart.query.filter_by(user_id=current_user.id).all()
    if cart_items:
        order = Order(user_id=current_user.id,
                    order_date=datetime.now(), status='Новый')
        db.session.add(order)

        for cart_item in cart_items:
            order_item = OrderItem(order_id=order.id,
                                product_id=cart_item.product_id, quantity=cart_item.quantity)
            db.session.add(order_item)

        db.session.commit()
        calculate_order_total(order.id)

        # Очистка корзины после оформления заказа
        for cart_item in cart_items:
            db.session.delete(cart_item)
        db.session.commit()

        flash('Заказ успешно оформлен!', 'success')
        return redirect(url_for('orders'))
    else:
        flash('Ваша корзина пуста.', 'warning')
        return redirect(url_for('index'))


@app.route('/order/<int:order_id>', methods=['GET', 'POST'])  # Добавьте методы 'POST'
@login_required
def order_details(order_id):
    order = Order.query.get_or_404(order_id)
    if order.user_id != current_user.id and not current_user.is_admin:
        flash('У вас нет доступа к этому заказу.', 'danger')
        return redirect(url_for('index'))

    if request.method == 'POST':
        message = request.form['message']
        if message:
            chat_message = ChatMessage(
                sender_id=current_user.id,
                receiver_id=order.user_id if current_user.is_manager else order.manager.user_id,
                order_id=order_id,
                message=message,
            )
            db.session.add(chat_message)
            db.session.commit()
            flash('Сообщение отправлено', 'success')
            return redirect(url_for('order_details', order_id=order_id)) 

    chat_messages = ChatMessage.query.filter_by(order_id=order_id).order_by(ChatMessage.timestamp.asc()).all()

    return render_template('order_details.html', order=order, messages=chat_messages)
# Маршруты для менеджера
@app.route('/manager')
@login_required
@manager_required
def manager_dashboard():
    return render_template('manager/manager_dashboard.html')

# Маршруты для товаров
@app.route('/manager/products')
@login_required
@manager_required
def manager_products():
    products = Product.query.all()
    return render_template('manager/manager_products.html', products=products)

@app.route('/manager/products/add', methods=['GET', 'POST'])
@login_required
@manager_required
def add_product():
    form = AddProductForm()
    if form.validate_on_submit():
        product = Product(
            name=form.name.data,
            description=form.description.data,
            price=form.price.data,
            category_id=form.category_id.data,
        )
        db.session.add(product)
        db.session.commit()  # Сначала сохраняем товар, чтобы получить его ID

        if form.image.data:
            product_folder = os.path.join(app.config['UPLOAD_FOLDER'], str(product.id))
            os.makedirs(product_folder, exist_ok=True)  # Создаем папку для товара
            filename = secure_filename(form.image.data.filename)
            file_path = os.path.join(product_folder, filename)
            form.image.data.save(file_path)
            product.image = filename  # Сохраняем только имя файла
            db.session.commit()

        flash('Товар успешно добавлен', 'success')
        return redirect(url_for('manager_products'))
    return render_template('manager/manager_add_product.html', form=form)
@app.route('/manager/products/<int:product_id>/edit', methods=['GET', 'POST'])
@login_required
@manager_required
def edit_product(product_id):
    product = Product.query.get_or_404(product_id)
    form = EditProductForm(obj=product)
    if form.validate_on_submit():
        if form.image.data:
            product_folder = os.path.join(app.config['UPLOAD_FOLDER'], str(product.id))
            os.makedirs(product_folder, exist_ok=True)  # Создаем папку, если ее нет
            filename = secure_filename(form.image.data.filename)
            file_path = os.path.join(product_folder, filename)
            form.image.data.save(file_path)
            product.image = filename
        product.name = form.name.data
        product.description = form.description.data
        product.price = form.price.data
        product.category_id = form.category_id.data
        db.session.commit()
        flash('Товар успешно изменен', 'success')
        return redirect(url_for('manager_products'))
    return render_template('manager/manager_edit_product.html', form=form, product=product)
@app.route('/manager/products/<int:product_id>/delete', methods=['POST'])
@login_required
@manager_required
def delete_product(product_id):
    product = Product.query.get_or_404(product_id)
    db.session.delete(product)
    db.session.commit()
    flash('Товар успешно удален', 'success')
    return redirect(url_for('manager/manager_products'))

# Маршруты для категорий
@app.route('/manager/categories')
@login_required
@manager_required
def manager_categories():
    categories = Category.query.all()
    return render_template('manager/manager_categories.html', categories=categories)

@app.route('/manager/categories/add', methods=['GET', 'POST'])
@login_required
@manager_required
def add_category():
    form = AddCategoryForm()
    if form.validate_on_submit():
        category = Category(name=form.name.data)  
        db.session.add(category)
        db.session.commit()
        flash('Категория успешно добавлена', 'success')
        return redirect(url_for('manager_categories'))
    return render_template('manager/manager_add_category.html', form=form)

@app.route('/manager/categories/<int:category_id>/edit', methods=['GET', 'POST'])
@login_required
@manager_required
def edit_category(category_id):
    category = Category.query.get_or_404(category_id)
    form = EditCategoryForm(obj=category)
    if form.validate_on_submit():
        category.name = form.name.data
        db.session.commit()
        flash('Категория успешно изменена', 'success')
        return redirect(url_for('manager_categories'))
    return render_template('manager/manager_edit_category.html', form=form, category=category)

@app.route('/manager/categories/<int:category_id>/delete', methods=['POST'])
@login_required
@manager_required
def delete_category(category_id):
    category = Category.query.get_or_404(category_id)
    db.session.delete(category)
    db.session.commit()
    flash('Категория успешно удалена', 'success')
    return redirect(url_for('manager/manager_categories'))


@app.route('/manager/orders')
@login_required
@manager_required
def manager_orders():
    orders = Order.query.all()
    return render_template('manager/manager_orders.html', orders=orders)

@app.route('/manager/orders/<int:order_id>', methods=['GET', 'POST']) # Добавьте POST
@login_required
@manager_required
def manager_order_details(order_id):
    """Просмотр деталей заказа."""
    order = Order.query.get_or_404(order_id)
    if request.method == 'POST':
        message = request.form['message']
        if message:
            chat_message = ChatMessage(
                sender_id=current_user.id,
                receiver_id=order.user_id,
                order_id=order_id,
                message=message,
            )
            db.session.add(chat_message)
            db.session.commit()
            flash('Сообщение отправлено', 'success')
            return redirect(url_for('manager_order_details', order_id=order_id)) 

    chat_messages = ChatMessage.query.filter_by(order_id=order_id).order_by(ChatMessage.timestamp.asc()).all()

    return render_template('manager/manager_order_details.html', order=order, messages=chat_messages)
@app.route('/manager/orders/<int:order_id>/update_status', methods=['POST'])
@login_required
@manager_required
def update_order_status(order_id):
    """Изменение статуса заказа."""
    order = Order.query.get_or_404(order_id)
    new_status = request.form.get('status')
    if new_status and new_status in ['Новый', 'В обработке', 'Отправлен', 'Доставлен', 'Отменен']:
        order.status = new_status
        db.session.commit()
        flash('Статус заказа успешно обновлен.', 'success')
    else:
        flash('Неверный статус заказа.', 'danger')
    return redirect(url_for('manager_order_details', order_id=order_id))

# ... ваш код ...

# ... ваш код ...

@app.route('/download_contract/<int:order_id>')
@login_required
def download_contract(order_id):
    order = Order.query.get_or_404(order_id)

    # Проверка прав доступа: только для менеджеров и владельца заказа
    if not current_user.is_manager and order.user_id != current_user.id:
        flash('У вас нет доступа к этому заказу.', 'danger')
        return redirect(url_for('index'))

    # Загрузка существующего шаблона
    template_path = os.path.join(app.root_path, 'static/contract_template.docx')
    doc = docx.Document(template_path)

    # Заполнение шаблона данными
    for paragraph in doc.paragraphs:
        if 'order_id' in paragraph.text:
            paragraph.text = paragraph.text.replace('order_id', str(order.id))
        if 'order_date' in paragraph.text:
            paragraph.text = paragraph.text.replace('order_date', order.order_date.strftime("%Y-%m-%d"))
        if 'customer_name' in paragraph.text:
            paragraph.text = paragraph.text.replace('customer_name', f'{order.user.first_name} {order.user.last_name}')
        if 'delivery_address' in paragraph.text:
            paragraph.text = paragraph.text.replace('delivery_address', order.address)
        if 'payment_method' in paragraph.text:
            paragraph.text = paragraph.text.replace('payment_method', order.payment_method)
        if 'total_price' in paragraph.text:
            paragraph.text = paragraph.text.replace('total_price', f'{order.total_price:.2f} руб.')

    # Заполнение информации о товарах
    items_text = ""
    for item in order.items:
        item_str = f"- {item.product.name} - Количество: {item.quantity}, Цена за единицу: {item.unit_price:.2f} руб., Сумма: {item.quantity * item.unit_price:.2f} руб.\n"
        items_text += item_str

    # Замените тег в документе
    for paragraph in doc.paragraphs:
        if '{items_text}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{items_text}', items_text)

    # Сохранение документа в память
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Возвращение документа как ответа
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'contract_{order.id}.docx'
    )

# ... ваш код ...

# ... ваш код ...
# Регистрация фильтра floatformat
# environment = Environment(app.jinja_env)
# environment.filters['floatformat'] = floatformat

if __name__ == '__main__':
    app.run(debug=True)